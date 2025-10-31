import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, PageBreak, Image,
    KeepTogether
)
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
from db.db import get_connection

# VISUELS
def save_pie_chart(labels, sizes, filename, title=""):
    import matplotlib.pyplot as plt
    plt.figure(figsize=(5,5))
    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    if title:
        plt.title(title)
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()
    
    
def save_bar_chart(x, y, filename, xlabel="", ylabel="", title=""):
    import matplotlib.pyplot as plt
    plt.figure(figsize=(7,4))
    plt.bar(x, y, color="#4682B4")
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

# RÉSUMÉ EXÉCUTIF AMÉLIORÉ
def generate_resume_executif(solde_cloture, evolution, total_recettes, total_depenses, nb_events, nb_membres):
    tendance = "positive" if solde_cloture > 0 else "négative"
    if evolution > 0:
        evo = f"en progression de {abs(evolution):.2f}€ par rapport à l'an passé"
    elif evolution < 0:
        evo = f"en baisse de {abs(evolution):.2f}€ par rapport à l'an passé"
    else:
        evo = "stable par rapport à l'an passé"
    part_recettes = f"Les recettes de l'exercice atteignent {total_recettes:.2f}€, traduisant la participation active de la communauté."
    part_depenses = f"Les dépenses, quant à elles, s'élèvent à {total_depenses:.2f}€, témoignant d'une gestion attentive."
    part_events = f"{nb_events} événement(s) ont été organisés pour impliquer les membres et soutenir les projets."
    part_membres = f"L'association compte {nb_membres} membres engagés."
    return (
        f"Le solde de clôture est de {solde_cloture:.2f}€, "
        f"ce qui traduit une situation financière {tendance}, {evo}.\n\n"
        f"{part_recettes}\n{part_depenses}\n{part_events}\n{part_membres}"
    )
    
    
# RÉCUPÉRATION ET STRUCTURATION DES DONNÉES
def get_event_details(c):
    events = c.execute("SELECT id, name, date FROM events ORDER BY date ASC").fetchall()
    details = []
    for ev in events:
        rec = c.execute("SELECT SUM(montant) as total FROM event_recettes WHERE event_id=?", (ev["id"],)).fetchone()
        dep = c.execute("SELECT SUM(montant) as total FROM event_depenses WHERE event_id=?", (ev["id"],)).fetchone()
        rec = rec["total"] or 0.0
        dep = dep["total"] or 0.0
        benef = rec - dep
        desc = f"L'événement {ev['name']} du {ev['date']} a réuni la communauté autour de moments conviviaux. Il a généré {rec:.2f} € de recettes pour {dep:.2f} € de dépenses, soit un bénéfice de {benef:.2f} €."
        details.append({
            "id": ev["id"],
            "name": ev["name"],
            "date": ev["date"],
            "recettes": rec,
            "depenses": dep,
            "benefice": benef,
            "description": desc,
            "commentaire": ""
        })
    return details
    
    
def get_data_for_bilan():
    from db.db import get_connection
    conn = get_connection()
    c = conn.cursor()
    cfg = c.execute("SELECT exercice, date, date_fin, solde_report, but_asso FROM config ORDER BY id DESC LIMIT 1").fetchone()
    exercice = cfg["exercice"] if cfg else "N/A"
    date = cfg["date"] if cfg else "N/A"
    date_fin = cfg["date_fin"] if cfg else "N/A"
    solde_ouverture = cfg["solde_report"] if cfg else 0.0
    but_asso = cfg["but_asso"] if cfg and "but_asso" in cfg.keys() and cfg["but_asso"] else "L'association a pour objet ... (à personnaliser)"
    comptes = c.execute("SELECT name, solde FROM comptes ORDER BY name").fetchall()
    comptes = [(cb["name"], cb["solde"]) for cb in comptes] if comptes else [("Banque", solde_ouverture)]
    membres = c.execute("SELECT name, prenom FROM membres ORDER BY name, prenom").fetchall()
    nb_membres = len(membres)
    liste_membres = [f"{m['name']} {m['prenom']}" for m in membres]
    event_details = get_event_details(c)
    nb_events = len(event_details)
    liste_events = [f"{e['name']} – {e['date']}" for e in event_details]
    subv = c.execute("SELECT SUM(montant) as total FROM dons_subventions").fetchone()
    subv_total = subv["total"] or 0.0
    rec_events = sum(ev["recettes"] for ev in event_details)
    recettes_par_categorie = []
    if subv_total: recettes_par_categorie.append(("Subventions", subv_total))
    if rec_events: recettes_par_categorie.append(("Recettes événements", rec_events))
    if not recettes_par_categorie:
        recettes_par_categorie = [("Aucune recette", 0.0)]
    total_recettes = sum(x[1] for x in recettes_par_categorie)
    dep_events = sum(ev["depenses"] for ev in event_details)
    dep_annexes = c.execute("SELECT SUM(montant) as total FROM depenses_regulieres").fetchone()
    dep_annexes_total = dep_annexes["total"] or 0.0
    dep_div = c.execute("SELECT SUM(montant) as total FROM depenses_diverses").fetchone()
    dep_div_total = dep_div["total"] or 0.0
    depenses_par_categorie = []
    if dep_events: depenses_par_categorie.append(("Dépenses événements", dep_events))
    if dep_annexes_total: depenses_par_categorie.append(("Frais annexes", dep_annexes_total))
    if dep_div_total: depenses_par_categorie.append(("Dépenses diverses", dep_div_total))
    if not depenses_par_categorie:
        depenses_par_categorie = [("Aucune dépense", 0.0)]
    total_depenses = sum(x[1] for x in depenses_par_categorie)
    retro = c.execute("SELECT SUM(montant) as total FROM retrocessions_ecoles").fetchone()
    retro_total = retro["total"] or 0.0
    retro_details = c.execute("SELECT date, ecole, montant, commentaire FROM retrocessions_ecoles ORDER BY date").fetchall()
    retro_details_list = [
        dict(date=r["date"], ecole=r["ecole"], montant=r["montant"], commentaire=r["commentaire"] or "")
        for r in retro_details
    ]
    solde_cloture = solde_ouverture + total_recettes - total_depenses
    solde_apres_retro = solde_cloture - retro_total
    treso = comptes
    prev = c.execute("SELECT solde_report, exercice FROM config WHERE id < (SELECT MAX(id) FROM config) ORDER BY id DESC LIMIT 1").fetchone()
    evolution = 0.0
    compo = []
    if prev:
        evolution = solde_cloture - prev["solde_report"]
        compo = [
            ["Solde ouverture", prev["solde_report"], solde_ouverture],
            ["Recettes totales", 0.0, total_recettes],
            ["Dépenses totales", 0.0, total_depenses],
            ["Solde clôture", prev["solde_report"], solde_cloture]
        ]
    recettes_majoritaires = ", ".join([cat for cat, montant in recettes_par_categorie if montant > 0 and cat != "Aucune recette"])
    depenses_majoritaires = ", ".join([cat for cat, montant in depenses_par_categorie if montant > 0 and cat != "Aucune dépense"])
    ratio = "N/A" if total_depenses == 0 else f"{total_recettes/total_depenses:.2f}"
    analyse = (
        f"Les recettes principales proviennent de : {recettes_majoritaires if recettes_majoritaires else 'aucune source majeure'}.\n"
        f"Les dépenses majeures sont : {depenses_majoritaires if depenses_majoritaires else 'aucune dépense majeure'}.\n"
        f"Le ratio recettes/dépenses est de {ratio}.\n"
        f"Nombre d'événements organisés : {nb_events}.\n"
        f"Nombre de membres engagés : {nb_membres}."
    )
    resume_executif = generate_resume_executif(solde_cloture, evolution, total_recettes, total_depenses, nb_events, nb_membres)
    conclusion = (
        "L'année écoulée a démontré la vitalité de l'association, soutenue par l'engagement bénévole et la diversité des actions menées. "
        "La gestion prudente des finances, la mobilisation des partenaires et la générosité des subventions ont permis de soutenir de nombreux projets scolaires."
    )
    data = {
        "exercice": f"{exercice} ({date} - {date_fin})",
        "date": date,
        "date_fin": date_fin,
        "resume_executif": resume_executif,
        "presentation": but_asso,
        "solde_ouverture": solde_ouverture,
        "total_recettes": total_recettes,
        "total_depenses": total_depenses,
        "solde_cloture": solde_cloture,
        "solde_apres_retro": solde_apres_retro,
        "recettes_par_categorie": recettes_par_categorie,
        "depenses_par_categorie": depenses_par_categorie,
        "analyse": analyse,
        "comparatif": compo,
        "tresorerie": treso,
        "conclusion": conclusion,
        # "etat_stock": etat_stock,   # SUPPRIMÉ
        "nb_membres": nb_membres,
        "liste_membres": liste_membres,
        "nb_events": nb_events,
        "liste_events": liste_events,
        "event_details": event_details,
        "subventions_total": subv_total,
        "retrocessions_total": retro_total,
        "retrocessions_details": retro_details_list,
    }
    conn.close()
    return data
    
    
# --- ÉDITION SECTION PAR SECTION (fenêtre personnalisée) ---
def edit_argumentaire_section(title, initial_text):
    import tkinter as tk
    result = {}
    win = tk.Toplevel()
    win.title(f"Édition – {title}")
    win.configure(bg="#f2f7fa")
    win.geometry("850x380")
    win.minsize(600, 220)
    tk.Label(win, text=title, font=("Arial", 15, "bold"), fg="#1b3a4b", bg="#f2f7fa").pack(pady=(18,10))
    txt = tk.Text(win, width=90, height=8, font=("Arial", 12), wrap="word", bg="#f6fafd", fg="#222")
    txt.pack(padx=16, pady=4, expand=True, fill="both")
    txt.insert("1.0", initial_text)
    txt.focus_set()
    txt.tag_add("sel", "1.0", "end")
    def valider():
        result["text"] = txt.get("1.0", "end-1c")
        win.destroy()
    btn = tk.Button(win, text="Valider / Suivant", font=("Arial", 12, "bold"), bg="#4c9ed9", fg="white", activebackground="#337ab7", activeforeground="white", width=18, height=2, command=valider)
    btn.pack(pady=14)
    win.grab_set()
    win.wait_window()
    return result.get("text", initial_text)
    
def edit_argumentaire(data):
    import datetime
    # Titre et sous-titre éditables
    data["titre"] = edit_argumentaire_section(
        "Titre principal du document", data.get("titre", "<b>Les Interactifs des Ecoles</b>")
    )
    data["sous_titre"] = edit_argumentaire_section(
        "Sous-titre (exercice)", data.get("sous_titre", f"Bilan financier argumenté<br/>Exercice : {data['exercice']}")
    )
    data["date_edition"] = edit_argumentaire_section(
        "Date d'édition du bilan", data.get("date_edition", f"Date d'édition du bilan : {datetime.date.today()}")
    )

    data["presentation"] = edit_argumentaire_section("Présentation / Historique", data.get("presentation", ""))
    data["resume_executif"] = edit_argumentaire_section("Résumé exécutif", data.get("resume_executif", ""))
    # Synthèse financière éditable (affichage tableau possible dans le texte)
    data["synthese_financiere"] = edit_argumentaire_section(
        "Synthèse financière (tableau)", data.get("synthese_financiere",
            f"Solde d'ouverture : {data['solde_ouverture']:.2f} €\nTotal recettes : {data['total_recettes']:.2f} €\nTotal dépenses : {data['total_depenses']:.2f} €\nSolde de clôture : {data['solde_cloture']:.2f} €"
        )
    )
    data["subventions"] = edit_argumentaire_section(
        "Subventions et partenaires", data.get("subventions", f"Total subventions versées : {data['subventions_total']:.2f} €")
    )
    data["retrocession"] = edit_argumentaire_section(
        "Rétrocession aux écoles", data.get("retrocession", f"Total rétrocédé : {data['retrocessions_total']:.2f} €")
    )
    data["analyse"] = edit_argumentaire_section("Analyse et commentaires", data.get("analyse", ""))
    data["conclusion"] = edit_argumentaire_section("Conclusion du/de la trésorier(e)", data.get("conclusion", ""))
    # Paragraphe solde après rétrocession
    data["solde_apres_retro"] = edit_argumentaire_section(
        "Solde restant après rétrocessions",
        data.get("solde_apres_retro", f"Solde restant : {data['solde_apres_retro']:.2f} €\nCe solde est conservé pour les frais à venir de l'association : frais bancaires, achats récurrents, trésorerie pour les prochains événements, etc.")
    )

    # Boucle sur les événements pour éditer description ET commentaire
    for ev in data.get("event_details", []):
        ev["description"] = edit_argumentaire_section(
            f"Événement : {ev['name']} ({ev['date']})", ev.get("description", "")
        )
        ev["commentaire"] = edit_argumentaire_section(
            f"Commentaire libre pour {ev['name']} ({ev['date']})", ev.get("commentaire", "")
        )

    try:
        conn = get_connection()
        c = conn.cursor()
        c.execute("UPDATE config SET but_asso=? WHERE id=(SELECT MAX(id) FROM config)", (data.get("presentation", ""),))
        conn.commit()
        conn.close()
    except Exception as e:
        print("Erreur lors de la sauvegarde du but_asso :", e)
    return True
    
def ask_date_cloture_if_missing(data):
    import tkinter as tk
    from tkinter import messagebox, simpledialog
    if not data.get("date_fin") or data["date_fin"] in (None, "", "None"):
        root = tk.Tk()
        root.withdraw()
        date_fin = None
        while not date_fin:
            date_fin = simpledialog.askstring(
                "Date de clôture manquante",
                "Veuillez saisir la date de clôture de l'exercice (format AAAA-MM-JJ) :"
            )
            if date_fin is None:
                messagebox.showwarning("Annulé", "Vous devez saisir une date de clôture pour générer le bilan.")
                return False
        data["date_fin"] = date_fin
        try:
            conn = get_connection()
            c = conn.cursor()
            c.execute("UPDATE config SET date_fin=? WHERE id=(SELECT MAX(id) FROM config)", (date_fin,))
            conn.commit()
            conn.close()
        except Exception as e:
            print("Erreur lors de la sauvegarde de la date de clôture :", e)
    return True
    
    
def export_bilan_argumente_pdf():
    root = tk.Tk()
    root.withdraw()
    filename = filedialog.asksaveasfilename(
        title="Enregistrer le bilan argumenté PDF",
        defaultextension=".pdf",
        filetypes=[("Fichier PDF", "*.pdf")],
        initialfile="Bilan_argumente.pdf"
    )
    if not filename:
        return

    data = get_data_for_bilan()
    if not ask_date_cloture_if_missing(data):
        return

    # Mise à jour de l'intitulé de l'exercice avec la bonne date de clôture
    exercice = data.get("exercice", "").split(" ")[0]
    date = data.get("date") or ""
    date_fin = data.get("date_fin") or ""
    data["exercice"] = f"{exercice} ({date} - {date_fin})"

    ok = edit_argumentaire(data)
    if not ok:
        messagebox.showinfo("Annulé", "Génération du bilan annulée.")
        return

    # Génération des camemberts
    if data["depenses_par_categorie"]:
        labels, sizes = zip(*data["depenses_par_categorie"])
        save_pie_chart(labels, sizes, "depenses_pie.png", "Dépenses")
    if data["recettes_par_categorie"]:
        labels2, sizes2 = zip(*data["recettes_par_categorie"])
        save_pie_chart(labels2, sizes2, "recettes_pie.png", "Recettes")

    # Détail des dépenses régulières et diverses
    conn = get_connection()
    c = conn.cursor()
    dep_reg = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_regulieres ORDER BY date_depense").fetchall()
    dep_div = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_diverses ORDER BY date_depense").fetchall()

    doc = SimpleDocTemplate(
        filename, pagesize=A4,
        rightMargin=38, leftMargin=38, topMargin=50, bottomMargin=38
    )
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("TitlePerso", parent=styles["Title"], fontSize=26, textColor="#184d72", alignment=1, spaceAfter=15)
    style_h1 = ParagraphStyle("H1Perso", parent=styles["Heading1"], fontSize=17, textColor="#20639b", alignment=1, spaceBefore=16, spaceAfter=8)
    style_h2 = ParagraphStyle("H2Perso", parent=styles["Heading2"], fontSize=14, textColor="#276c93", alignment=1, spaceBefore=10, spaceAfter=6)
    style_text = ParagraphStyle("NormalPerso", parent=styles["Normal"], fontSize=11, textColor="#222", alignment=1, leading=15, spaceAfter=4)
    style_italic = ParagraphStyle("ItalicPerso", parent=styles["Italic"], fontSize=11, textColor="#3a3a3a", alignment=1, leading=13, italic=True)

    elements = []
    # TITRE & INFOS
    elements.append(KeepTogether([
        Paragraph(data['titre'], style_title),
        Spacer(1, 12),
        Paragraph(data['sous_titre'], style_h1),
        Spacer(1, 12),
        Paragraph(data['date_edition'], style_text),
        PageBreak(),
    ]))

    elements.append(KeepTogether([
        Paragraph("Présentation / Historique de l'association", style_h1),
        Paragraph(data['presentation'], style_text),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Membres & Bénévoles", style_h2),
        Paragraph(f"Nombre total de membres : {data['nb_membres']}", style_text),
        Table([["Name Prénom"]]+[[m] for m in data["liste_membres"]], style=[
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.lightcyan]),
        ]) if data["liste_membres"] else Spacer(1, 1),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Résumé exécutif", style_h1),
        Paragraph(data['resume_executif'], style_text),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Synthèse financière", style_h1),
        Paragraph(data['synthese_financiere'], style_text),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Subventions et partenaires", style_h2),
        Paragraph(data['subventions'], style_text),
        Spacer(1, 14),
    ]))

    retro_block = [
        Paragraph("Rétrocession aux écoles", style_h2),
        Paragraph(data['retrocession'], style_text),
    ]
    if data["retrocessions_details"]:
        table_retro = [["Date", "École", "Montant (€)", "Commentaire"]]
        for r in data["retrocessions_details"]:
            table_retro.append([
                r["date"],
                r["ecole"],
                f"{r['montant']:.2f}",
                r["commentaire"] or ""
            ])
        retro_block.append(Table(table_retro, style=[
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('ROWBACKGROUNDS', (0,1),(-1,-1), [colors.whitesmoke, colors.lightcyan])
        ]))
        retro_block.append(Spacer(1, 8))
    retro_block += [
        Paragraph("Solde restant après rétrocessions", style_h2),
        Paragraph(data['solde_apres_retro'], style_text),
        Spacer(1, 14)
    ]
    elements.append(KeepTogether(retro_block))

    elements.append(KeepTogether([
        Paragraph("Détail des recettes", style_h2),
        Table(
            [["Catégorie", "Montant"]] + [[cat, f"{montant:.2f} €"] for cat, montant in data['recettes_par_categorie']],
            style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.lightcyan, colors.whitesmoke])]
        ),
        Spacer(1, 8),
        Image("recettes_pie.png", width=170, height=170) if data["recettes_par_categorie"] else Spacer(1, 1),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Détail des dépenses", style_h2),
        Table(
            [["Catégorie", "Montant"]] + [[cat, f"{montant:.2f} €"] for cat, montant in data['depenses_par_categorie']],
            style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.lightcyan])]
        ),
        Spacer(1, 8),
        Image("depenses_pie.png", width=170, height=170) if data["depenses_par_categorie"] else Spacer(1, 1),
        Spacer(1, 14),
    ]))

    elements.append(KeepTogether([
        Paragraph("Dépenses régulières", style_h2),
        Table(
            [["Date", "Catégorie", "Montant (€)", "Commentaire"]]+
            [[d["date_depense"], d["categorie"], f"{d['montant']:.2f}", d["commentaire"] or ""] for d in dep_reg],
            style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.lightcyan])]
        ) if dep_reg else Paragraph("Aucune dépense régulière enregistrée.", style_text),
        Spacer(1, 14),
    ]))
    elements.append(KeepTogether([
        Paragraph("Dépenses diverses", style_h2),
        Table(
            [["Date", "Catégorie", "Montant (€)", "Commentaire"]]+
            [[d["date_depense"], d["categorie"], f"{d['montant']:.2f}", d["commentaire"] or ""] for d in dep_div],
            style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.lightcyan, colors.whitesmoke])]
        ) if dep_div else Paragraph("Aucune dépense diverse enregistrée.", style_text),
        Spacer(1, 14),
    ]))

    elements.append(Paragraph("Détail des événements", style_h1))
    for ev in data["event_details"]:
        ev_block = [
            Paragraph(f"{ev['name']} - {ev['date']}", style_h2),
            Paragraph(ev["description"], style_text),
            Paragraph(ev["commentaire"], style_italic) if ev["commentaire"] else Spacer(1, 1),
            Table([
                ["Recettes", f"{ev['recettes']:.2f} €"],
                ["Dépenses", f"{ev['depenses']:.2f} €"],
                ["Bénéfice", f"{ev['benefice']:.2f} €"]
            ], style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.lightcyan])])
        ]
        if ev["recettes"] or ev["depenses"]:
            save_pie_chart(["Recettes", "Dépenses"], [ev["recettes"], ev["depenses"]], f"event_{ev['id']}_pie.png", ev["name"])
            ev_block.append(Image(f"event_{ev['id']}_pie.png", width=110, height=110))
        ev_block.append(Spacer(1, 14))
        elements.append(KeepTogether(ev_block))

    elements.append(KeepTogether([
        Paragraph("Analyse et commentaires", style_h1),
        Paragraph(data['analyse'], style_text),
        Spacer(1, 14),
    ]))

    if data["comparatif"]:
        compo_block = [
            Paragraph("Comparaison avec l'exercice précédent", style_h2),
            Table(
                [["Poste", "Année N-1", "Année N"]]
                + [[poste, f"{n_1:.2f} €", f"{n:.2f} €"] for poste, n_1, n in data['comparatif']],
                style=[('ALIGN', (0,0), (-1,-1), 'CENTER'), ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.lightcyan])]
            ),
            Spacer(1, 14),
        ]
        elements.append(KeepTogether(compo_block))

    elements.append(KeepTogether([
        Paragraph("Conclusion du/de la trésorier(e)", style_h1),
        Paragraph(data['conclusion'], style_text),
        Spacer(1, 20),
    ]))

    try:
        doc.build(elements)
        messagebox.showinfo("Succès", f"Bilan argumenté généré : {filename}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors de la génération : {e}")
        
        
# --- VERSION ENTIÈREMENT ÉDITABLE DE export_bilan_argumente_word ---
def export_bilan_argumente_word():
    root = tk.Tk()
    root.withdraw()
    filename = filedialog.asksaveasfilename(
        title="Enregistrer le bilan argumenté Word",
        defaultextension=".docx",
        filetypes=[("Fichier Word", "*.docx")],
        initialfile="Bilan_argumente.docx"
    )
    if not filename:
        return
    data = get_data_for_bilan()
    ok = edit_argumentaire(data)
    if not ok:
        messagebox.showinfo("Annulé", "Génération du bilan annulée.")
        return

    # Génération des camemberts
    if data["depenses_par_categorie"]:
        labels, sizes = zip(*data["depenses_par_categorie"])
        save_pie_chart(labels, sizes, "depenses_pie.png", "Dépenses")
    if data["recettes_par_categorie"]:
        labels2, sizes2 = zip(*data["recettes_par_categorie"])
        save_pie_chart(labels2, sizes2, "recettes_pie.png", "Recettes")

    # Récupérer les détails des dépenses régulières et diverses
    conn = get_connection()
    c = conn.cursor()
    dep_reg = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_regulieres ORDER BY date_depense").fetchall()
    dep_div = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_diverses ORDER BY date_depense").fetchall()

    doc = Document()
    doc.add_heading(data['titre'], 0)
    doc.add_heading(data['sous_titre'], 1)
    doc.add_paragraph(data['date_edition'])
    doc.add_page_break()

    doc.add_heading("Présentation / Historique de l'association", level=1)
    doc.add_paragraph(data['presentation'])

    doc.add_heading("Membres & Bénévoles", level=2)
    doc.add_paragraph(f"Nombre total de membres : {data['nb_membres']}")
    if data["liste_membres"]:
        tM = doc.add_table(rows=1, cols=1)
        tM.rows[0].cells[0].text = "Name Prénom"
        for m in data["liste_membres"]:
            tM.add_row().cells[0].text = m

    doc.add_heading("Résumé exécutif", level=1)
    doc.add_paragraph(data['resume_executif'])

    doc.add_heading("Synthèse financière", level=1)
    doc.add_paragraph(data['synthese_financiere'])

    doc.add_heading("Subventions et partenaires", level=2)
    doc.add_paragraph(data['subventions'])

    doc.add_heading("Rétrocession aux écoles", level=2)
    doc.add_paragraph(data['retrocession'])
    if data["retrocessions_details"]:
        tR = doc.add_table(rows=1, cols=4)
        tR.rows[0].cells[0].text = "Date"
        tR.rows[0].cells[1].text = "École"
        tR.rows[0].cells[2].text = "Montant (€)"
        tR.rows[0].cells[3].text = "Commentaire"
        for r in data["retrocessions_details"]:
            row = tR.add_row().cells
            row[0].text = r["date"]
            row[1].text = r["ecole"]
            row[2].text = f"{r['montant']:.2f}"
            row[3].text = r["commentaire"] or ""
    doc.add_paragraph("")
    doc.add_heading("Solde restant après rétrocessions", level=2)
    doc.add_paragraph(data['solde_apres_retro'])

    doc.add_heading("Détail des recettes", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Catégorie"
    t.rows[0].cells[1].text = "Montant"
    for cat, montant in data['recettes_par_categorie']:
        row = t.add_row().cells
        row[0].text = cat
        row[1].text = f"{montant:.2f} €"
    if data["recettes_par_categorie"]:
        doc.add_paragraph("Diagramme des recettes :")
        doc.add_picture("recettes_pie.png", width=Inches(2.5))

    doc.add_heading("Détail des dépenses", level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "Catégorie"
    t2.rows[0].cells[1].text = "Montant"
    for cat, montant in data['depenses_par_categorie']:
        row = t2.add_row().cells
        row[0].text = cat
        row[1].text = f"{montant:.2f} €"
    if data["depenses_par_categorie"]:
        doc.add_paragraph("Diagramme des dépenses :")
        doc.add_picture("depenses_pie.png", width=Inches(2.5))

    # Détail des dépenses régulières
    doc.add_heading("Dépenses régulières", level=2)
    if dep_reg:
        tdepreg = doc.add_table(rows=1, cols=4)
        tdepreg.rows[0].cells[0].text = "Date"
        tdepreg.rows[0].cells[1].text = "Catégorie"
        tdepreg.rows[0].cells[2].text = "Montant (€)"
        tdepreg.rows[0].cells[3].text = "Commentaire"
        for d in dep_reg:
            row = tdepreg.add_row().cells
            row[0].text = d["date_depense"] or ""
            row[1].text = d["categorie"] or ""
            row[2].text = f"{d['montant']:.2f}"
            row[3].text = d["commentaire"] or ""
    else:
        doc.add_paragraph("Aucune dépense régulière enregistrée.")

    # Détail des dépenses diverses
    doc.add_heading("Dépenses diverses", level=2)
    if dep_div:
        tdepdiv = doc.add_table(rows=1, cols=4)
        tdepdiv.rows[0].cells[0].text = "Date"
        tdepdiv.rows[0].cells[1].text = "Catégorie"
        tdepdiv.rows[0].cells[2].text = "Montant (€)"
        tdepdiv.rows[0].cells[3].text = "Commentaire"
        for d in dep_div:
            row = tdepdiv.add_row().cells
            row[0].text = d["date_depense"] or ""
            row[1].text = d["categorie"] or ""
            row[2].text = f"{d['montant']:.2f}"
            row[3].text = d["commentaire"] or ""
    else:
        doc.add_paragraph("Aucune dépense diverse enregistrée.")

    doc.add_heading("Détail des événements", level=1)
    for ev in data["event_details"]:
        doc.add_heading(f"{ev['name']} - {ev['date']}", level=2)
        doc.add_paragraph(ev["description"])
        if ev["commentaire"]:
            doc.add_paragraph(ev["commentaire"], style="Intense Quote")
        table_ev = doc.add_table(rows=3, cols=2)
        table_ev.rows[0].cells[0].text = "Recettes"
        table_ev.rows[0].cells[1].text = f"{ev['recettes']:.2f} €"
        table_ev.rows[1].cells[0].text = "Dépenses"
        table_ev.rows[1].cells[1].text = f"{ev['depenses']:.2f} €"
        table_ev.rows[2].cells[0].text = "Bénéfice"
        table_ev.rows[2].cells[1].text = f"{ev['benefice']:.2f} €"
        if ev["recettes"] or ev["depenses"]:
            save_pie_chart(["Recettes", "Dépenses"], [ev["recettes"], ev["depenses"]], f"event_{ev['id']}_pie.png", ev["name"])
            doc.add_picture(f"event_{ev['id']}_pie.png", width=Inches(1.5))
        doc.add_paragraph("")

    doc.add_heading("Analyse et commentaires", level=1)
    doc.add_paragraph(data['analyse'])

    if data["comparatif"]:
        doc.add_heading("Comparaison avec l'exercice précédent", level=2)
        t3 = doc.add_table(rows=1, cols=3)
        t3.rows[0].cells[0].text = "Poste"
        t3.rows[0].cells[1].text = "Année N-1"
        t3.rows[0].cells[2].text = "Année N"
        for poste, n_1, n in data['comparatif']:
            row = t3.add_row().cells
            row[0].text = poste
            row[1].text = f"{n_1:.2f} €"
            row[2].text = f"{n:.2f} €"

    doc.add_heading("Conclusion du/de la trésorier(e)", level=1)
    doc.add_paragraph(data['conclusion'])

    try:
        doc.save(filename)
        messagebox.showinfo("Succès", f"Bilan argumenté généré : {filename}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors de la génération : {e}")
    root = tk.Tk()
    root.withdraw()
    filename = filedialog.asksaveasfilename(
        title="Enregistrer le bilan argumenté Word",
        defaultextension=".docx",
        filetypes=[("Fichier Word", "*.docx")],
        initialfile="Bilan_argumente.docx"
    )
    if not filename:
        return
    data = get_data_for_bilan()
    ok = edit_argumentaire(data)
    if not ok:
        messagebox.showinfo("Annulé", "Génération du bilan annulée.")
        return

    # Génération des camemberts
    if data["depenses_par_categorie"]:
        labels, sizes = zip(*data["depenses_par_categorie"])
        save_pie_chart(labels, sizes, "depenses_pie.png", "Dépenses")
    if data["recettes_par_categorie"]:
        labels2, sizes2 = zip(*data["recettes_par_categorie"])
        save_pie_chart(labels2, sizes2, "recettes_pie.png", "Recettes")

    # Récupérer les détails des dépenses régulières et diverses
    conn = get_connection()
    c = conn.cursor()
    dep_reg = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_regulieres ORDER BY date_depense").fetchall()
    dep_div = c.execute("SELECT date_depense, categorie, montant, commentaire FROM depenses_diverses ORDER BY date_depense").fetchall()

    doc = Document()
    doc.add_heading("Les Interactifs des Ecoles", 0)
    doc.add_heading(f"Bilan financier argumenté - Exercice : {data['exercice']}", 1)
    doc.add_paragraph(f"Date d'édition du bilan : {datetime.date.today()}")
    doc.add_page_break()

    doc.add_heading("Présentation / Historique de l'association", level=1)
    doc.add_paragraph(data['presentation'])

    doc.add_heading("Membres & Bénévoles", level=2)
    doc.add_paragraph(f"Nombre total de membres : {data['nb_membres']}")
    if data["liste_membres"]:
        tM = doc.add_table(rows=1, cols=1)
        tM.rows[0].cells[0].text = "Name Prénom"
        for m in data["liste_membres"]:
            tM.add_row().cells[0].text = m

    doc.add_heading("Résumé exécutif", level=1)
    doc.add_paragraph(data['resume_executif'])

    doc.add_heading("Synthèse financière", level=1)
    tab = doc.add_table(rows=1, cols=2)
    for i, champ in enumerate(["Solde d'ouverture", "Total recettes", "Total dépenses", "Solde de clôture"]):
        row_cells = tab.add_row().cells if i else tab.rows[0].cells
        row_cells[0].text = champ
        val = [data['solde_ouverture'], data['total_recettes'], data['total_depenses'], data['solde_cloture']][i]
        row_cells[1].text = f"{val:.2f} €"

    doc.add_heading("Subventions et partenaires", level=2)
    doc.add_paragraph(f"Total subventions versées : {data['subventions_total']:.2f} €")

    doc.add_heading("Rétrocession aux écoles", level=2)
    doc.add_paragraph(f"Total rétrocédé : {data['retrocessions_total']:.2f} €")
    if data["retrocessions_details"]:
        tR = doc.add_table(rows=1, cols=4)
        tR.rows[0].cells[0].text = "Date"
        tR.rows[0].cells[1].text = "École"
        tR.rows[0].cells[2].text = "Montant (€)"
        tR.rows[0].cells[3].text = "Commentaire"
        for r in data["retrocessions_details"]:
            row = tR.add_row().cells
            row[0].text = r["date"]
            row[1].text = r["ecole"]
            row[2].text = f"{r['montant']:.2f}"
            row[3].text = r["commentaire"] or ""
    doc.add_paragraph("")
    doc.add_heading("Solde restant après rétrocessions", level=2)
    doc.add_paragraph(f"Solde restant : {data['solde_apres_retro']:.2f} €")
    doc.add_paragraph(
        "Ce solde est conservé pour les frais à venir de l'association : frais bancaires, achats récurrents, trésorerie pour les prochains événements, etc."
    )

    doc.add_heading("Détail des recettes", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Catégorie"
    t.rows[0].cells[1].text = "Montant"
    for cat, montant in data['recettes_par_categorie']:
        row = t.add_row().cells
        row[0].text = cat
        row[1].text = f"{montant:.2f} €"
    if data["recettes_par_categorie"]:
        doc.add_paragraph("Diagramme des recettes :")
        doc.add_picture("recettes_pie.png", width=Inches(2.5))

    doc.add_heading("Détail des dépenses", level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "Catégorie"
    t2.rows[0].cells[1].text = "Montant"
    for cat, montant in data['depenses_par_categorie']:
        row = t2.add_row().cells
        row[0].text = cat
        row[1].text = f"{montant:.2f} €"
    if data["depenses_par_categorie"]:
        doc.add_paragraph("Diagramme des dépenses :")
        doc.add_picture("depenses_pie.png", width=Inches(2.5))

    # Détail des dépenses régulières
    doc.add_heading("Dépenses régulières", level=2)
    if dep_reg:
        tdepreg = doc.add_table(rows=1, cols=4)
        tdepreg.rows[0].cells[0].text = "Date"
        tdepreg.rows[0].cells[1].text = "Catégorie"
        tdepreg.rows[0].cells[2].text = "Montant (€)"
        tdepreg.rows[0].cells[3].text = "Commentaire"
        for d in dep_reg:
            row = tdepreg.add_row().cells
            row[0].text = d["date_depense"] or ""
            row[1].text = d["categorie"] or ""
            row[2].text = f"{d['montant']:.2f}"
            row[3].text = d["commentaire"] or ""
    else:
        doc.add_paragraph("Aucune dépense régulière enregistrée.")

    # Détail des dépenses diverses
    doc.add_heading("Dépenses diverses", level=2)
    if dep_div:
        tdepdiv = doc.add_table(rows=1, cols=4)
        tdepdiv.rows[0].cells[0].text = "Date"
        tdepdiv.rows[0].cells[1].text = "Catégorie"
        tdepdiv.rows[0].cells[2].text = "Montant (€)"
        tdepdiv.rows[0].cells[3].text = "Commentaire"
        for d in dep_div:
            row = tdepdiv.add_row().cells
            row[0].text = d["date_depense"] or ""
            row[1].text = d["categorie"] or ""
            row[2].text = f"{d['montant']:.2f}"
            row[3].text = d["commentaire"] or ""
    else:
        doc.add_paragraph("Aucune dépense diverse enregistrée.")

    doc.add_heading("Détail des événements", level=1)
    for ev in data["event_details"]:
        doc.add_heading(f"{ev['name']} - {ev['date']}", level=2)
        doc.add_paragraph(ev["description"])
        if ev["commentaire"]:
            doc.add_paragraph(ev["commentaire"], style="Intense Quote")
        table_ev = doc.add_table(rows=3, cols=2)
        table_ev.rows[0].cells[0].text = "Recettes"
        table_ev.rows[0].cells[1].text = f"{ev['recettes']:.2f} €"
        table_ev.rows[1].cells[0].text = "Dépenses"
        table_ev.rows[1].cells[1].text = f"{ev['depenses']:.2f} €"
        table_ev.rows[2].cells[0].text = "Bénéfice"
        table_ev.rows[2].cells[1].text = f"{ev['benefice']:.2f} €"
        if ev["recettes"] or ev["depenses"]:
            save_pie_chart(["Recettes", "Dépenses"], [ev["recettes"], ev["depenses"]], f"event_{ev['id']}_pie.png", ev["name"])
            doc.add_picture(f"event_{ev['id']}_pie.png", width=Inches(1.5))
        doc.add_paragraph("")

    doc.add_heading("Analyse et commentaires", level=1)
    doc.add_paragraph(data['analyse'])

    if data["comparatif"]:
        doc.add_heading("Comparaison avec l'exercice précédent", level=2)
        t3 = doc.add_table(rows=1, cols=3)
        t3.rows[0].cells[0].text = "Poste"
        t3.rows[0].cells[1].text = "Année N-1"
        t3.rows[0].cells[2].text = "Année N"
        for poste, n_1, n in data['comparatif']:
            row = t3.add_row().cells
            row[0].text = poste
            row[1].text = f"{n_1:.2f} €"
            row[2].text = f"{n:.2f} €"

    doc.add_heading("Conclusion du/de la trésorier(e)", level=1)
    doc.add_paragraph(data['conclusion'])

    try:
        doc.save(filename)
        messagebox.showinfo("Succès", f"Bilan argumenté généré : {filename}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors de la génération : {e}")